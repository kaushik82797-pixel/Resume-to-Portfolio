"""
resume_reader.py
----------------
Extracts plain text content from PDF, DOCX, and TXT resume files.
Supports automatic format detection based on file extension and robust error handling.
"""

import os
from typing import Optional


class ResumeReaderError(Exception):
    """Custom exception raised when resume reading fails."""
    pass


def read_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF (pymupdf)."""
    try:
        import pymupdf as fitz
    except ImportError:
        try:
            import fitz
        except ImportError:
            raise ResumeReaderError(
                "PyMuPDF library is missing. Please run: pip install pymupdf"
            )

    try:
        doc = fitz.open(file_path)
        text_content = []
        for page in doc:
            text_content.append(page.get_text())
        doc.close()
        extracted = "\n".join(text_content).strip()
        if not extracted:
            raise ResumeReaderError(f"PDF file '{file_path}' appears to be empty or image-only.")
        return extracted
    except Exception as e:
        if isinstance(e, ResumeReaderError):
            raise e
        raise ResumeReaderError(f"Failed to read PDF file '{file_path}': {str(e)}")


def read_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
    except ImportError:
        raise ResumeReaderError(
            "python-docx library is missing. Please run: pip install python-docx"
        )

    try:
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract text from tables if present
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_content = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_content:
                    table_text.append(" | ".join(row_content))
        
        full_text = "\n".join(paragraphs + table_text).strip()
        if not full_text:
            raise ResumeReaderError(f"DOCX file '{file_path}' contains no readable text.")
        return full_text
    except Exception as e:
        if isinstance(e, ResumeReaderError):
            raise e
        raise ResumeReaderError(f"Failed to read DOCX file '{file_path}': {str(e)}")


def read_txt(file_path: str) -> str:
    """Extract text from a TXT file using native Python file operations."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read().strip()
        if not content:
            raise ResumeReaderError(f"TXT file '{file_path}' is empty.")
        return content
    except UnicodeDecodeError:
        # Fallback encoding try
        try:
            with open(file_path, "r", encoding="latin-1") as f:
                content = f.read().strip()
            if not content:
                raise ResumeReaderError(f"TXT file '{file_path}' is empty.")
            return content
        except Exception as e:
            raise ResumeReaderError(f"Failed to read TXT file with fallback encoding: {str(e)}")
    except Exception as e:
        if isinstance(e, ResumeReaderError):
            raise e
        raise ResumeReaderError(f"Failed to read TXT file '{file_path}': {str(e)}")


def extract_resume_text(file_path: str) -> str:
    """
    Main entry point for extracting text from a resume file.
    Automatically detects format based on extension.
    """
    if not os.path.exists(file_path):
        raise ResumeReaderError(f"File not found: '{file_path}'. Please check the path and try again.")

    if not os.path.isfile(file_path):
        raise ResumeReaderError(f"Path '{file_path}' is a directory, not a valid resume file.")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return read_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        if ext == ".doc":
            raise ResumeReaderError("Older '.doc' format is not directly supported. Please convert it to '.docx' or '.pdf'.")
        return read_docx(file_path)
    elif ext in [".txt", ".md"]:
        return read_txt(file_path)
    else:
        raise ResumeReaderError(
            f"Unsupported file format '{ext}'. Supported extensions are: .pdf, .docx, .txt"
        )
